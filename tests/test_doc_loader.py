# tests/test_doc_loader.py
import os
import tempfile
from scraper.doc_loader import load_docx, load_xlsx, load_all_docs
from knowledge.db import get_all_knowledge

def test_load_docx():
    from docx import Document
    doc = Document()
    doc.add_paragraph("Грант выдаётся студентам с отличным средним баллом от 4.5")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        fname = f.name
    try:
        chunks = load_docx(fname)
        assert len(chunks) > 0
        assert any("Грант" in c for c in chunks)
    finally:
        os.remove(fname)

def test_load_xlsx():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["Специальность", "Код", "Форма обучения"])
    ws.append(["Программирование", "09.02.07", "Очная"])
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        wb.save(f.name)
        fname = f.name
    try:
        chunks = load_xlsx(fname)
        assert any("Программирование" in c for c in chunks)
    finally:
        os.remove(fname)

async def test_load_all_docs_adds_to_db():
    from docx import Document
    doc = Document()
    doc.add_paragraph("Специальность Информационные системы в организации — очная форма")
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "spec.docx")
        doc.save(path)
        await load_all_docs(tmpdir)
    rows = await get_all_knowledge()
    assert len(rows) > 0
